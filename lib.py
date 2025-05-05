import os
from datetime import datetime

import requests

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from dotenv import load_dotenv


load_dotenv()

application_id = os.getenv("APPLICATION_ID")
base_url = os.getenv("BASE_URL")
stats_url = os.getenv("STATS_URL")

nations = {
    "commonwealth": "Commonwealth",
    "europe": "European",
    "france": "French",
    "germany": "German",
    "italy": "Italian",
    "japan": "Japanese",
    "netherlands": "Dutch",
    "pan_america": "Pan-American",
    "pan_asia": "Pan-Asian",
    "spain": "Spanish",
    "uk": "British",
    "usa": "American",
    "ussr": "Soviet",
}

ship_types = {
    "AirCarrier": "Aircraft Carrier",
    "Battleship": "Battleship",
    "Cruiser": "Cruiser",
    "Destroyer": "Destroyer",
    "Submarine": "Submarine",
}

real_steel_ids = [4292753392, 3768465392, 4266538992, 4274927600, 4289607664, 3764271088, 4288559088, 4264441840,
                  4286461936, 3762173936, 4272830448, 3761125360, 3750639600, 3551410160, 3540924400, 4281219056,
                  3760076784, 3655219184, 4293834736, 4292786160, 4291737584, 4290689008, 3767449584, 4289640432,
                  4248745968, 4247697392, 4183734256, 4288591856, 4182685680, 4077828080, 3763255280, 3668883440,
                  3553540080, 4181637104, 4076779504, 3762206704, 3730749424, 3720263664, 3760109552, 4273911792,
                  4074682352, 3550394352, 4293867504, 4290721776, 4281284592, 4288624624, 3765385200, 3555670000,
                  4259264496, 3764336624, 3763288048, 4286527472, 3742316528, 3553572848, 4282333168, 3762239472,
                  3751753712, 3542038512, 4276041712, 3761190896, 3529455600, 4185896944, 4267685872, 4183799792,
                  4265588720, 4181702640, 3751786480, 3741300720, 3730814960, 4273977328, 4179605488, 3760175088,
                  4078352368, 4076255216, 4074158064, 3759585264, 3549870064,
                  
                  4293801680, 4280170192, 4292753104, 4269684432, 4291704528, 4290655952, 4268635856, 4267587280,
                  4184749776, 4183701200, 4078843600, 3554555600, 4182652624, 4077795024, 3763222224, 4181604048,
                  4076746448, 3751687888, 4075697872, 3540924112, 4282267344, 4256085712, 4258182864, 4279154384,
                  3767449296, 4290688720, 4281251536, 4289640144, 4079924944, 3765352144, 4287542992, 4286494416,
                  3752769232, 4285445840, 4276008656, 4255037136, 3540956880, 3655251664, 4283381456, 4293867216,
                  4287575760, 4288624336, 3764336336, 3743364816, 4284430032, 3752802000, 3761190608, 4276041424,
                  4185896656, 4183799504, 4181702352, 3751786192, 3550459600, 3751196368,
                  
                  4186846672, 4185798096, 4293801424, 4078843344, 4077794768, 3763221968, 3752736208, 3865982416,
                  3551409616, 4292785616, 4284397008, 4187928016, 4293834192, 4186879440, 3767449040, 3756963280,
                  4185830864, 4291737040, 3765351888, 3744380368, 3764303312, 3659445712, 4181636560, 3762206160,
                  3709777360, 4185863632, 3765384656, 3753850320, 4183209424, 4181112272, 3761681872, 4179015120,
                  
                  4187895600, 4186847024, 4184749872, 4182652720, 4077795120, 3762173744, 4181604144, 3751687984,
                  4292785968, 3768497968, 4186879792, 4185831216, 4184782640, 4183734064, 3764303664, 3753817904,
                  4181636912, 3762206512, 3530471216, 4186912560, 4082054960, 3767482160, 4081006384, 4079957808,
                  3763287856, 3543086896, 4292818736, 4181669680, 4078352176, 4076255024, 4074157872,
                  
                  4187895760, 4186847184, 3767416784, 4185798608, 4184750032, 4183701456, 3764271056, 4182652880,
                  3543021520, 4181604304, 3752736720, 4180555728, 4179507152, 4188977104, 4187928528, 4186879952,
                  4185831376, 4184782800, 4079925200, 3765352400, 4183734224, 4078876624, 3753818064, 3743332304,
                  3544102864, 4182685648, 3763255248, 4181637072, 3751720912, 3741235152, 4186912720, 4082055120,
                  3767482320, 4081006544, 4079957968, 3765385168, 4292818896, 4183766992, 4078909392, 3721344976,
                  3743365072, 4182718416, 3763288016, 3752802256, 3742316496, 3543087056, 3762239440, 4081039312,
                  4183799760, 3751786448, 4181702608, 3762272208, 3741300688, 3542071248, 4074747856, 4178556880,
                  4183209936, 4181112784, 3761682384, 4179015632, 3549870032,
                  
                  3767416176, 3763221872, 3752736112, 3760076144, 4293834096, 4281251184, 4280202608, 4279154032,
                  4278105456, 3764303216, 3753817456, 4277056880,
                  
                  4187895632, 4185798480, 4184749904, 4279121744, 3766368080, 4278073168, 4183701328, 3764270928,
                  4277024592, 4182652752, 4181604176, 3762173776, 4180555600, 4284364624, 4188976976, 4187928400,
                  4185831248, 4184782672, 4183734096, 3753817936, 3544102736, 4182685520, 3760109392, 4185864016,
                  4184815440, 3764336464, 3763287888, 4181669712, 3751753552, 3764369232,
                  
                  4187895536, 4186846960, 4185798384, 4184749808, 4183701232, 3764270832, 4182652656, 4179506928,
                  4188976880, 4186879728, 4185831152, 4184782576, 4183734000, 3764303600, 4182685424, 3763255024,
                  3658397424, 4185863920, 4184815344, 3765384944, 4183766768, 4181669616, 3762239216,
                  
                  4188976560, 4187927984, 4186879408, 4185830832, 4184782256, 4183733680, 3764303280,
                  
                  3761125136, 3655218960, 4188976912, 4187928336, 4186879760, 4185831184, 3762206480,
                  
                  4187895088, 4186846512, 4185797936, 4184749360, 3555603760, 4079891760, 4078843184, 4077794608,
                  3769513264, 3543020848, 4181603632, 3762173232, 4180555056, 3761124656, 3760076080, 4179506480,
                  3655218480, 4188976432, 3764303152, 3766433072,
                  
                  4188976464, 4187927888, 4184782160, 4182685008, 3763254608, 3762206032,
                  
                  4186846416, 4183700688, 3764270288, 4182652112, 4181603536, 4076745936, 3762173136, 3751687376,
                  4180554960, 4179506384, 4188976336, 3767448784, 4184782032, 3762205904, 3543577808]

def download_image_if_needed(url, save_dir="images"):
    """
    Downloads an image from the given URL only if it's not already downloaded.
    Returns the local path to the image.
    """
    os.makedirs(save_dir, exist_ok=True)
    filename = os.path.basename(url.split("?")[0])  # Strip URL parameters
    local_path = os.path.join(save_dir, filename)

    if not os.path.exists(local_path):
        response = requests.get(url)
        if response.status_code == 200:
            with open(local_path, "wb") as f:
                f.write(response.content)
        else:
            raise Exception(f"Failed to download image: {url} (status {response.status_code})")

    return local_path

def extract_ship_info(ship: dict) -> dict:
    ship_id = ship["id"]
    ship_nation = ship["nation"]
    ship_tier = ship["tier"]
    ship_type = ship["type"]
    ship_name = ship["name"]
    ship_description = ship["description"]
    ship_image = ship["image"]
    is_premium = ship["is_premium"]
    is_special = ship["is_special"]
    has_demo_profile = ship["has_demo_profile"]
    is_real_steel = True if int(ship_id) in real_steel_ids else False
    
    image_path = download_image_if_needed(ship_image) if len(ship_image) > 0 else None
    
    if len(ship_description) > 0:
        ship_description = ship_description.replace('\n', ' ')
    
    prefix = (
        "Premium " if is_premium else
        "Special " if is_special else
        "Demo " if has_demo_profile else
        ""
    )
    type_str = f"{prefix}{ship_types.get(ship_type, ship_type)}"
    nation_str = nations.get(ship_nation, ship_nation)
    ship_tier_roman = int_to_roman(ship_tier)
    
    return {
        "id": ship_id,
        "nation": ship_nation,
        "nation_str": nation_str,
        "tier": ship_tier,
        "tier_roman": ship_tier_roman,
        "type": ship_type,
        "type_str": type_str,
        "name": ship_name,
        "description": ship_description,
        "image": image_path,
        "is_premium": is_premium,
        "is_special": is_special,
        "has_demo_profile": has_demo_profile,
        "is_real_steel": is_real_steel,
    }

def get_usable_width(doc):
    """Returns usable width of the page in cm (page width minus left/right margins)"""
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    usable_width_cm = (page_width - left_margin - right_margin) / Cm(1)
    return usable_width_cm

def int_to_roman(num):
    roman_numerals = {
        1000: "M", 900: "CM", 500: "D", 400: "CD",
        100: "C", 90: "XC", 50: "L", 40: "XL",
        10: "X", 9: "IX", 5: "V", 4: "IV", 1: "I"
    }
    roman_str = ""
    for val, symbol in roman_numerals.items():
        while num >= val:
            roman_str += symbol
            num -= val
    return roman_str

def is_enclosed_in_brackets(string):
    return string.startswith("[") and string.endswith("]")

def is_azure_lane(string):
    return string.startswith("AL ")

def is_arpeggio_of_blue_steel(string):
    return string.startswith("ARP ")

def is_black_friday(string):
    return string.endswith(" B")

def is_color_ship(string):
    return string.endswith(" CLR")

def is_victory_lap_ship(string):
    return string.endswith(" VL")

def is_supertest_ship(string):
    return string.endswith(" ST")

def load_expected_values_from_api():
    res = requests.get(stats_url)
    ships = res.json().get("data", {})
    result = {}
    for ship_id, ship_stats in ships.items():
        if not isinstance(ship_stats, dict):
            continue
        result[ship_id] = {
            "dmg": ship_stats.get("average_damage_dealt", 0),
            "frg": ship_stats.get("average_frags", 0),
            "wr": ship_stats.get("win_rate", 0),
        }
    return result

def load_ships_from_api(al: bool = True, arp: bool = True, bf: bool = True, clr: bool = True, vl: bool = True,
                        st: bool = True, as_list: bool = False, sort_keys: list = None) -> dict | list:
    current_page = 1
    max_page = 9999
    
    result = {}
    
    while current_page <= max_page:
        page_url = f"{base_url}/encyclopedia/ships/?application_id={application_id}&page_no={current_page}"
        res = requests.get(page_url)
        data = res.json()
        max_page = min(max_page, data["meta"]["page_total"])
        for key, value in data["data"].items():
            ship_id = key
            ship_nation = value["nation"]
            ship_tier = value["tier"]
            ship_type = value["type"]
            ship_name = value["name"].replace("\xa0", " ")
            ship_description = value["description"].replace("\xa0", " ")
            is_premium = value["is_premium"]
            is_special = value["is_special"]
            has_demo_profile = value["has_demo_profile"]
            ship_image = value.get("images", {}).get("large")
            if is_enclosed_in_brackets(ship_name):
                continue
            if not al and is_azure_lane(ship_name):
                continue
            if not arp and is_arpeggio_of_blue_steel(ship_name):
                continue
            if not bf and is_black_friday(ship_name):
                continue
            if not clr and is_color_ship(ship_name):
                continue
            if not vl and is_victory_lap_ship(ship_name):
                continue
            if not st and is_supertest_ship(ship_name):
                continue
            
            result[ship_id] = {
                "id": ship_id,
                "nation": ship_nation,
                "nation_str": nations.get(ship_nation, ship_nation),
                "tier": ship_tier,
                "tier_roman": int_to_roman(ship_tier),
                "type": ship_type,
                "type_str": ship_types.get(ship_type, ship_type),
                "name": ship_name,
                "description": ship_description,
                "image": ship_image,
                "is_premium": is_premium,
                "is_special": is_special,
                "has_demo_profile": has_demo_profile,
            }
        
        current_page += 1
    
    if as_list:
        return ship_dict_to_list(result, sort_keys=sort_keys)
    
    return result

def load_stats_from_api(account_id: int) -> dict:
    url = f"{base_url}/ships/stats/?application_id={application_id}&account_id={account_id}"
    res = requests.get(url)
    result = {}
    for ship_stats in res.json().get("data", {}).get(f"{account_id}", []):
        ship_id = ship_stats.get("ship_id")
        
        total_battles = ship_stats.get("battles", 0)
        last_battle_time = ship_stats.get("last_battle_time")
        distance = ship_stats.get("distance", 0)
        
        pvp = ship_stats.get("pvp", {})
        battles = pvp.get("battles", 0)
        art_agro = pvp.get("art_agro", 0)
        capture_points = pvp.get("capture_points", 0)
        damage_dealt = pvp.get("damage_dealt", 0)
        damage_scouting = pvp.get("damage_scouting", 0)
        dropped_capture_points = pvp.get("dropped_capture_points", 0)
        frags = pvp.get("frags", 0)
        max_damage_dealt = pvp.get("max_damage_dealt", 0)
        max_damage_scouting = pvp.get("max_damage_scouting", 0)
        max_frags_battle = pvp.get("max_frags_battle", 0)
        max_planes_killed = pvp.get("max_planes_killed", 0)
        max_ships_spotted = pvp.get("max_ships_spotted", 0)
        max_total_agro = pvp.get("max_total_agro", 0)
        max_xp = pvp.get("max_xp", 0)
        planes_killed = pvp.get("planes_killed", 0)
        ships_spotted = pvp.get("ships_spotted", 0)
        survived_battles = pvp.get("survived_battles", 0)
        team_capture_points = pvp.get("team_capture_points", 0)
        team_dropped_capture_points = pvp.get("team_dropped_capture_points", 0)
        torpedo_agro = pvp.get("torpedo_agro", 0)
        wins = pvp.get("wins", 0)
        xp = pvp.get("xp", 0)
        
        result[ship_id] = {
            "total_battles": total_battles,
            "battles": battles,
            "wins": wins,
            "survived": survived_battles,
            "damage": damage_dealt,
            "max_damage": max_damage_dealt,
            "kills": frags,
            "max_kills": max_frags_battle,
            "spotted": ships_spotted,
            "max_spotted": max_ships_spotted,
            "spotting_damage": damage_scouting,
            "max_spotting_damage": max_damage_scouting,
            "planes_killed": planes_killed,
            "max_planes_killed": max_planes_killed,
            "xp": xp,
            "max_xp": max_xp,
            "agro": art_agro + torpedo_agro,
            "max_agro": max_total_agro,
            "last_battle_time": last_battle_time,
            "distance": distance,
            "capture_points": capture_points,
            "dropped_capture_points": dropped_capture_points,
            "team_capture_points": team_capture_points,
            "team_dropped_capture_points": team_dropped_capture_points,
        }
    
    return result
    
def remove_first_paragraph(doc):
    if doc.paragraphs:
        first_para = doc.paragraphs[0]
        if not first_para.text.strip():
            p = first_para._element
            p.getparent().remove(p)
            p._p = p._element = None

def set_cell_background(cell, color_hex):
    """
    Set background color of a cell.
    :param cell: The cell object
    :param color_hex: Hex color code without '#' (e.g., 'FF0000' for red)
    """
    # Get the cell's XML
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=None, start=None, bottom=None, end=None):
    """
    Set cell padding (margins) individually.
    :param end: padding at the right side of the cell (twips)
    :param bottom: padding at the bottom of the cell (twips)
    :param start: padding at the left side of the cell (twips)
    :param top: padding at the top of the cell (twips)
    :param cell: cell to set margins for
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    for direction, margin in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        if margin is not None:
            node = tcMar.find(qn(f'w:{direction}'))
            if node is None:
                node = OxmlElement(f'w:{direction}')
                tcMar.append(node)
            node.set(qn('w:w'), str(margin))
            node.set(qn('w:type'), 'dxa')  # dxa = twips

def ship_dict_to_list(ships: dict, sort_keys: list = None) -> list:
    result = list(ships.values())
    
    if not sort_keys:
        return result
    
    def sort_key_func(ship):
        key_list = []
        for item in sort_keys:
            if isinstance(item, str):
                # Default ascending
                key_list.append(ship.get(item))
            elif isinstance(item, (tuple, list)) and len(item) == 2:
                field, ascending = item
                value = ship.get(field)
                # For descending fields, invert the value for correct sorting
                if isinstance(value, (int, float)):
                    key_list.append(value if ascending else -value)
                elif isinstance(value, str):
                    key_list.append(value if ascending else ''.join(chr(255 - ord(c)) for c in value))
                elif isinstance(value, datetime):
                    key_list.append(value if ascending else datetime.min - (value - datetime.min))
                else:
                    key_list.append(value)  # fallback
            else:
                raise ValueError("sort_keys must be a list of strings or (field, ascending) tuples")
        return tuple(key_list)
    
    return sorted(result, key=sort_key_func)

def write_cell(cell, text, bold=False, italic=False, align=None, background_color=None, padding=None, font_color=None):
    """
    Write text into a Word table cell, with optional bold, italic, alignment, background color, padding, and font color.
    """
    cell._element.clear_content()
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic

    if font_color is not None:
        run.font.color.rgb = font_color

    if align is not None:
        paragraph.alignment = align

    if background_color is not None:
        set_cell_background(cell, background_color)

    if padding is not None:
        set_cell_margins(cell, **padding)
