#!/usr/bin/env python3
from docx import Document

from lib import extract_ship_info, load_ships_from_api, remove_first_paragraph


current_page = 1
max_page = 9999

doc = Document('./input/template2.docx')
remove_first_paragraph(doc)

ships_sorted = load_ships_from_api(as_list=True, sort_keys=["name", "nation"])

current_letter = None

for ship in ships_sorted:
    data = extract_ship_info(ship)
    
    first_letter = data["name"][0].upper()
    
    if current_letter != first_letter:
        doc.add_paragraph(f"{first_letter}", style="Heading 1")
    current_letter = first_letter
    
    para = doc.add_paragraph("")
    name_run = para.add_run(f"{data['name']} ")
    tier_run = para.add_run(f"{data['tier_roman']}", style="Subtle Emphasis")
    if data["is_premium"]:
        premium_run = para.add_run("*", style="Emphasis")
    elif data["is_special"]:
        special_run = para.add_run("**", style="Emphasis")
    tab_run = para.add_run("\t")
    id_run = para.add_run(f"{data["id"]}", style="Intense Emphasis")
    
doc.save('./output/ids.docx')
