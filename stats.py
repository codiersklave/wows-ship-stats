#!/usr/bin/env python3
import argparse
import time

from datetime import datetime, timedelta

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

from lib import (load_expected_values_from_api, load_ships_from_api, load_stats_from_api, remove_first_paragraph,
                 ship_dict_to_list, write_cell)


def get_ship_nation(nation_str: str, adjective: bool = False) -> str:
    if nation_str == "CW": return "Commonwealth"
    if nation_str == "EU": return "European" if adjective else "Europe"
    if nation_str == "FR": return "French" if adjective else "France"
    if nation_str == "DE": return "German" if adjective else "Germany"
    if nation_str == "IT": return "Italian" if adjective else "Italy"
    if nation_str == "JP": return "Japanese" if adjective else "Japan"
    if nation_str == "NL": return "Dutch" if adjective else "Netherlands"
    if nation_str == "AM": return "Pan-American" if adjective else "Pan-America"
    if nation_str == "AS": return "Pan-Asian" if adjective else "Pan-Asia"
    if nation_str == "ES": return "Spanish" if adjective else "Spain"
    if nation_str == "UK": return "British" if adjective else "United Kingdom"
    if nation_str == "US": return "American" if adjective else "United States"
    if nation_str == "SU": return "Soviet" if adjective else "Soviet Union"
    
    return ""

def get_ship_type(type_str: str) -> str:
    if type_str == "A": return "Aircraft Carriers"
    if type_str == "B": return "Battleships"
    if type_str == "C": return "Cruisers"
    if type_str == "D": return "Destroyers"
    if type_str == "S": return "Submarines"
    
    return ""


ship_type_abbreviations = {
    "AirCarrier": "A",
    "Battleship": "B",
    "Cruiser": "C",
    "Destroyer": "D",
    "Submarine": "S",
}

nation_codes = {
    "commonwealth": "CW",
    "europe": "EU",
    "france": "FR",
    "germany": "DE",
    "italy": "IT",
    "japan": "JP",
    "netherlands": "NL",
    "pan_america": "AM",
    "pan_asia": "AS",
    "spain": "ES",
    "uk": "UK",
    "usa": "US",
    "ussr": "SU",
}


parser = argparse.ArgumentParser(description="Generate a docx file with stats for a given account.")
parser.add_argument("account_id", type=int, help="Account ID")
parser.add_argument("--days", type=int, default=30, help="Number of days to include in the docx file.")
parser.add_argument("--type", type=str, choices=["all", "A", "B", "C", "D", "S"], default="all", help="Filter by ship type.")
parser.add_argument("--nation", type=str, choices=["all", "CW", "EU", "FR", "DE", "IT", "JP", "NL", "AM", "AS", "ES", "UK", "US", "SU"], default="all", help="Filter by nation.")
parser.add_argument("--order", type=str, choices=["date", "name"], default="date", help="Order by date or name.")
parser.add_argument("--ship", type=str, help="The ship ID we are interested in. If specified, only this ship will be included in the docx file.")
args = parser.parse_args()

account_id = args.account_id

alternate_cell_bg = "efefef"
cell_vpadding = {"top": 72, "bottom": 72}

ships = load_ships_from_api()
stats = load_stats_from_api(account_id)
expected = load_expected_values_from_api()
result = {}

for ship_id, ship_info in ships.items():
    ship_stats = stats.get(int(ship_id), {})
    if not ship_stats.get("total_battles", None):
        continue
    
    result[ship_id] = ship_info | ship_stats

if args.order == "date":
    sort_keys = [("last_battle_time", False), "name"]
else:
    sort_keys = ["name"]

sorted_list = ship_dict_to_list(result, sort_keys=sort_keys)

now = datetime.now()
now_ts = int(time.time())

doc = Document('./input/template.docx')
remove_first_paragraph(doc)

if args.type == "all" and args.nation == "all":
    title_str = "All Ships Types"
elif args.type != "all" and args.nation == "all":
    title_str = f"{get_ship_type(args.type)}"
elif args.type == "all" and args.nation != "all":
    title_str = f"{get_ship_nation(args.nation)}"
else:
    title_str = f"{get_ship_nation(args.nation, True)} {get_ship_type(args.type)}"

title_str += f" (Last {args.days} days)"

title_str += f" - {now.strftime('%d %B %Y, %H:%M')}"

doc.core_properties.title = title_str
doc.core_properties.author = f"Player {account_id}"

for data in sorted_list:
    if args.type != "all" and ship_type_abbreviations.get(data["type"], "") != args.type:
        continue
    
    if args.nation != "all" and nation_codes.get(data["nation"], "") != args.nation:
        continue
    
    if args.ship and args.ship != data["id"]:
        continue
    
    battles = data["battles"]
    wins = data["wins"]
    damage = data["damage"]
    kills = data["kills"]
    spotted = data["spotted"]
    spotting_damage = data["spotting_damage"]
    planes_killed = data["planes_killed"]
    xp = data["xp"]
    
    if battles == 0:
        continue
    
    last_battle_ts = data["last_battle_time"]
    last_battle_dt = datetime.fromtimestamp(last_battle_ts)
    last_battle_str = last_battle_dt.strftime("%d %B %Y, %H:%M")
    
    if now - last_battle_dt > timedelta(days=args.days):
        continue
    
    win_rate = wins / battles * 100 if battles > 0 else 0
    avg_dmg = damage / battles if battles > 0 else 0
    avg_frg = kills / battles if battles > 0 else 0
    avg_spotted = spotted / battles if battles > 0 else 0
    avg_spotting_dmg = spotting_damage / battles if battles > 0 else 0
    avg_planes_killed = planes_killed / battles if battles > 0 else 0
    avg_xp = xp / battles if battles > 0 else 0
    
    expected_values = expected.get(data["id"])
    pr = None
    
    if expected_values:
        exp_dmg = expected_values.get("dmg")
        exp_frg = expected_values.get("frg")
        exp_wr = expected_values.get("wr")
        
        if exp_dmg and exp_frg and exp_wr:
            r_dmg = avg_dmg / exp_dmg
            r_frg = avg_frg / exp_frg
            r_wr = win_rate / exp_wr
            
            n_dmg = max(0, (r_dmg - 0.4) / (1 - 0.4))
            n_frg = max(0, (r_frg - 0.1) / (1 - 0.1))
            n_wr = max(0, (r_wr - 0.7) / (1 - 0.7))
            
            pr = 700 * n_dmg + 300 * n_frg + 150 * n_wr
    
    ship_info = f"(Tier {data["tier_roman"]} {data["nation_str"]} "
    if data["is_premium"]: ship_info += "Premium "
    elif data["is_special"]: ship_info += "Special "
    ship_info += f"{data['type_str']})"
    
    para = doc.add_paragraph(f"{data["name"]} ", style="Heading 1")
    para.paragraph_format.keep_with_next = True
    run_descr = para.add_run(f"{ship_info}   ")
    run_descr.style = "Subtle Emphasis"
    run_pr = para.add_run(f"PR: {pr:,.0f}" if pr else "PR: n/a")
    run_pr.style = "Emphasis"
    run_id = para.add_run(f"\t{data["id"]}")
    run_id.style = "Intense Emphasis"
    
    anchor = doc.add_paragraph()
    anchor.paragraph_format.keep_with_next = True
    
    table = doc.add_table(rows=7, cols=15)
    # table.style = "MyTableGrid"
    
    deep_red = RGBColor(0xC0, 0x00, 0x00)
    
    # Row 1
    
    write_cell(table.cell(0, 0).merge(table.cell(0, 1)), text="Battles", bold=True, padding=cell_vpadding)
    write_cell(table.cell(0, 2).merge(table.cell(0, 4)), text=f"{data['battles']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    write_cell(table.cell(0, 5).merge(table.cell(0, 6)), text="Victories", bold=True, padding=cell_vpadding)
    write_cell(table.cell(0, 7).merge(table.cell(0, 9)), text=f"{data['wins']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    write_cell(table.cell(0, 10).merge(table.cell(0, 11)), text="Win Rate", bold=True, padding=cell_vpadding)
    write_cell(table.cell(0, 12).merge(table.cell(0, 14)), text=f"{win_rate:.2f}%", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    # Row 2
    
    write_cell(table.cell(1, 0).merge(table.cell(1, 1)), text="Damage", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(1, 2).merge(table.cell(1, 4)), text=f"{data['damage']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    write_cell(table.cell(1, 5).merge(table.cell(1, 6)), text="Average", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(1, 7).merge(table.cell(1, 9)), text=f"{avg_dmg:,.2f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    write_cell(table.cell(1, 10).merge(table.cell(1, 11)), text="Maximum", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(1, 12).merge(table.cell(1, 14)), text=f"{data['max_damage']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    # Row 3
    
    write_cell(table.cell(2, 0).merge(table.cell(2, 1)), text="Kills", bold=True, padding=cell_vpadding)
    write_cell(table.cell(2, 2).merge(table.cell(2, 4)), text=f"{data['kills']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    write_cell(table.cell(2, 5).merge(table.cell(2, 6)), text="Average", bold=True, padding=cell_vpadding)
    write_cell(table.cell(2, 7).merge(table.cell(2, 9)), text=f"{avg_frg:,.2f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    write_cell(table.cell(2, 10).merge(table.cell(2, 11)), text="Maximum", bold=True, padding=cell_vpadding)
    write_cell(table.cell(2, 12).merge(table.cell(2, 14)), text=f"{data['max_kills']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    # Row 4
    
    write_cell(table.cell(3, 0).merge(table.cell(3, 1)), text="Spotted", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(3, 2).merge(table.cell(3, 4)), text=f"{data['spotted']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    write_cell(table.cell(3, 5).merge(table.cell(3, 6)), text="Average", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(3, 7).merge(table.cell(3, 9)), text=f"{avg_spotted:,.2f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    write_cell(table.cell(3, 10).merge(table.cell(3, 11)), text="Maximum", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(3, 12).merge(table.cell(3, 14)), text=f"{data['max_spotted']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    # Row 5
    
    write_cell(table.cell(4, 0).merge(table.cell(4, 1)), text="Assist", bold=True, padding=cell_vpadding)
    write_cell(table.cell(4, 2).merge(table.cell(4, 4)), text=f"{data['spotting_damage']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    write_cell(table.cell(4, 5).merge(table.cell(4, 6)), text="Average", bold=True, padding=cell_vpadding)
    write_cell(table.cell(4, 7).merge(table.cell(4, 9)), text=f"{avg_spotting_dmg:,.2f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    write_cell(table.cell(4, 10).merge(table.cell(4, 11)), text="Maximum", bold=True, padding=cell_vpadding)
    write_cell(table.cell(4, 12).merge(table.cell(4, 14)), text=f"{data['max_spotting_damage']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding)
    
    # Row 6
    
    write_cell(table.cell(5, 0).merge(table.cell(5, 1)), text="Planes", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(5, 2).merge(table.cell(5, 4)), text=f"{data['planes_killed']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    write_cell(table.cell(5, 5).merge(table.cell(5, 6)), text="Average", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(5, 7).merge(table.cell(5, 9)), text=f"{avg_planes_killed:,.2f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    write_cell(table.cell(5, 10).merge(table.cell(5, 11)), text="Maximum", bold=True, background_color=alternate_cell_bg, padding=cell_vpadding)
    write_cell(table.cell(5, 12).merge(table.cell(5, 14)), text=f"{data['max_planes_killed']:,.0f}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT, background_color=alternate_cell_bg, padding=cell_vpadding)
    
    # Row 7
    
    write_cell(table.cell(6, 0).merge(table.cell(6, 1)), text="XP", bold=True, padding=cell_vpadding)
    write_cell(table.cell(6, 2).merge(table.cell(6, 4)), text=f"{data['xp']:,.0f}", bold=True, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding, font_color=deep_red)
    
    write_cell(table.cell(6, 5).merge(table.cell(6, 6)), text="Average", bold=True, padding=cell_vpadding)
    write_cell(table.cell(6, 7).merge(table.cell(6, 9)), text=f"{avg_xp:,.2f}", bold=True, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding, font_color=deep_red)
    
    write_cell(table.cell(6, 10).merge(table.cell(6, 11)), text="Maximum", bold=True, padding=cell_vpadding)
    write_cell(table.cell(6, 12).merge(table.cell(6, 14)), text=f"{data['max_xp']:,.0f}", bold=True, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, padding=cell_vpadding, font_color=deep_red)
    
    doc.add_paragraph("", "Normal")
    para = doc.add_paragraph("")
    run_label = para.add_run("Last Battle: ")
    run_label.bold = True
    run_value = para.add_run(f"{last_battle_str}.")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    if now - last_battle_dt > timedelta(days=365):
        # Set run color to red
        run_value.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red
    elif now - last_battle_dt > timedelta(days=90):
        run_value.font.color.rgb = RGBColor(0xFF, 0x6F, 0x42)
    else:
        run_value.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    
    '''
    win_rate = data["wins"] / data["battles"] * 100 if data["battles"] > 0 else 0
    write_cell(table.cell(0, 3).merge(table.cell(0, 4)), "Win Rate:", True)
    write_cell(table.cell(0, 5), f"{win_rate:.2f}%", False, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    
    write_cell(table.cell(0, 6).merge(table.cell(0, 7)), "Survived:", True)
    write_cell(table.cell(0, 8), f"{data['survived']:,.0f}", False, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    
    avg_dmg = data["damage"] / data["battles"] if data["battles"] > 0 else 0
    write_cell(table.cell(1, 0).merge(table.cell(1, 1)), "Total Damage:", True)
    write_cell(table.cell(1, 2), f"{data['damage']:,.0f}", False, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    write_cell(table.cell(1, 3).merge(table.cell(1, 4)), "Avg. Damage", True)
    write_cell(table.cell(1, 5), f"{avg_dmg:,.0f}", False, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    write_cell(table.cell(1, 6).merge(table.cell(1, 7)), "Max. Damage", True)
    write_cell(table.cell(1, 8), f"{data['max_damage']:,.0f}", False, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    '''
    
    doc.add_paragraph("", style="No Spacing")
    doc.add_paragraph("", style="Normal")

doc.save(f"./output/stats_{account_id}_{now_ts}.docx")
