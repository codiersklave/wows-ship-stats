#!/usr/bin/env python3
from docx import Document
from docx.shared import Cm

from lib import extract_ship_info, load_ships_from_api, remove_first_paragraph, get_usable_width

include_azure_lane = False
include_arpeggio_of_blue_steel = False
include_black_friday = False
include_color_ships = False
include_victory_lap_ships = False
include_supertest_ships = False

current_page = 1
max_page = 9999

doc = Document('./input/template.docx')
remove_first_paragraph(doc)

ships_sorted = load_ships_from_api(al=include_azure_lane, arp=include_arpeggio_of_blue_steel, bf=include_black_friday,
                                   clr=include_color_ships, vl=include_victory_lap_ships, st=include_supertest_ships,
                                   as_list=True, sort_keys=["name", "nation"])

for ship in ships_sorted:
    data = extract_ship_info(ship)
    
    ship_name_str = data["name"]
    if data["is_real_steel"]:
        ship_name_str += "*"
    
    para = doc.add_paragraph(f"{ship_name_str} ", style="Heading 1")
    run_descr = para.add_run(
        f"(Tier {data["tier_roman"]} {data["nation_str"]} {data["type_str"]})\t")
    run_descr.style = "Subtle Emphasis"
    run_id = para.add_run(f"{data["id"]}")
    run_id.style = "Intense Emphasis"
    
    if data["image"]:
        full_width_cm = get_usable_width(doc)
        doc.add_picture(data["image"], width=Cm(full_width_cm))
    
    doc.add_paragraph(f"{data["description"]}")
    doc.add_paragraph("", style="No Spacing")
    
    doc.add_paragraph("", style="Normal")

doc.save('./output/ships.docx')
